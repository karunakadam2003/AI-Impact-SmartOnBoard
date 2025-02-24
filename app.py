from dotenv import load_dotenv
import os
import json
import docx
import google.generativeai as genai
import re
from unstructured.partition.auto import partition

load_dotenv()

# Configure Gemini API
GOOGLE_API_KEY = os.getenv("GENAI_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-pro')  # Or use 'gemini-pro' or other available models

# def read_word_document(file_path):
#     """Reads text from a Word document.  Uses python-docx library."""
#     try:
#         doc = docx.Document(file_path)
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
#         return '\n'.join(full_text)
#     except Exception as e:
#         print(f"Error reading document with docx: {e}")
#         return None

# def read_word_document_unstructured(file_path):
#     """Reads text from a Word document using the unstructured library."""
#     try:
#         elements = partition(filename=file_path)
#         full_text = "\n".join([str(el.text) for el in elements])
#         return full_text
#     except Exception as e:
#         print(f"Error reading document with unstructured: {e}")
#         return None

# def extract_key_value_pairs(text, prompt_template):
#     """Extracts key-value pairs using Gemini API."""
#     prompt = prompt_template.format(text=text)

#     try:
#         response = model.generate_content(prompt)
#         print("i extracted key value pairs")
#         return response.text
#     except Exception as e:
#         print(f"Error during Gemini API call: {e}")
#         return None


# def post_process_kv_pairs_cot(gemini_output):
#     """
#     Post-processes the Gemini output with Chain of Thought (CoT) to extract clean key-value pairs.
#     Handles the reasoning steps and extracted values separately.
#     This is a complex function and will likely need adjustments based on Gemini's actual output.
#     """
#     key_value_pairs = {}
#     lines = gemini_output.split("\n")
#     current_key = None
#     reasoning = []

#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue  # Skip empty lines

#         if line.startswith("Client Legal Name:"):
#             current_key = "Client Legal Name"
#             reasoning = []  # Reset reasoning for the new key
#         elif line.startswith("Doing Business As :"):
#             current_key = "Doing Business As"
#             reasoning = []
#         elif line.startswith("INdustry :"):
#             current_key = "Industry"
#             reasoning = []  # Note: Corrected "INdustry"
#         elif line.startswith("Date of Incorporation :"):
#             current_key = "Date of Incorporation"
#             reasoning = []
#         elif line.startswith("Client ID:"):
#             current_key = "Client ID"
#             reasoning = []
#         elif line.startswith("Primary Contact Name:"):
#             current_key = "Primary Contact Name"
#             reasoning = []
#         elif line.startswith("Primary Contact Email:"):
#             current_key = "Primary Contact Email"
#             reasoning = []
#         elif line.startswith("Number of Employees:"):
#             current_key = "Number of Employees"
#             reasoning = []
#         elif line.startswith("Annual Revenue:"):
#             current_key = "Annual Revenue"
#             reasoning = []
#         elif line.startswith("UBOs (Ultimate Beneficial Owners):"):
#             current_key = "UBOs (Ultimate Beneficial Owners)"
#             reasoning = []
#         elif line.startswith("Primary Bank (Currently):"):
#             current_key = "Primary Bank (Currently)"
#             reasoning = []
#         elif line.startswith("Services Requested:"):
#             current_key = "Services Requested"
#             reasoning = []

#         elif line.startswith("Reasoning:"):
#             # Append the reasoning to the reasoning list
#             reasoning.append(line[len("Reasoning:"):].strip())  # Remove "Reasoning:" prefix

#         elif line.startswith("Extracted Value:"):
#             extracted_value = line[len("Extracted Value:"):].strip()  # Remove "Extracted Value:" prefix
#             if current_key:
#                 key_value_pairs[current_key] = {
#                     "value": extracted_value,
#                     "reasoning": "\n".join(reasoning),  # Join reasoning steps into a single string
#                 }
#                 current_key = None  # Reset current_key after extracting the value
#             else:
#                 print(f"Warning: Extracted value found without a corresponding key: {line}")

#     return key_value_pairs
# def save_kv_pairs_to_json(kv_pairs, json_path):
#     """Saves key-value pairs to a JSON file."""
#     try:
#         with open(json_path, 'w') as json_file:
#             json.dump(kv_pairs, json_file, indent=4)
#         print(f"Key-value pairs saved to {json_path}")
#     except Exception as e:
#         print(f"Error saving key-value pairs to JSON: {e}")



# if __name__ == "__main__":
#     file_path = "sample.docx"  # Replace with your Word document path

#     # 1. Read the Document
#     # document_text = read_word_document(file_path)  # using python-docx

#     document_text = read_word_document_unstructured(file_path)  # Using unstructured

#     if document_text:
#         # 2. Define the Prompt (Crucial)
#         prompt_template = """
#         You are an expert financial onboarding specialist. Your task is to extract key information from a commercial client onboarding document.

#         Here's the document content you will be reviewing:

#         Text:
#         {text}

#         Based on the document content provided, extract the following information. For each item:

#         1.  Explain your reasoning process step-by-step.
#         2.  Provide the extracted value.
#         3.  If a value is not found or cannot be reliably determined from the provided text, explicitly state "Value not found".

#         Specifically, extract the following:

#         -   Client Legal Name: The full legal name of the commercial client.
#             Reasoning: Identify the phrase labeled "Client Legal Name" or a similar designation. Extract the text that follows.
#             Extracted Value:

#         -   Doing Business As : The DBA name of the commercial client, if different from the legal name.
#             Reasoning: Identify the section labeled "Doing Business As (DBA)" or a similar designation. Extract the text that follows.
#             Extracted Value:

#         -   INdustry : The industry sector in which the client operates.
#             Reasoning: Identify the section labeled "Industry" or a similar designation. Extract the text that follows.
#             Extracted Value:

#         -   Date of Incorporation : The date of incorporation for the client's business.
#             Reasoning: Identify the section labeled "Date of Incorporation" or a similar designation. Extract the date provided.
#             Extracted Value:

#         -   Client ID: The unique client identification number assigned by the bank.
#             Reasoning: Identify the section labeled "Client ID" or a similar designation. Extract the unique identification number.
#             Extracted Value:

#         -   Primary Contact Name: The full name of the primary contact person.
#             Reasoning: Locate the section labeled "Primary Contact." Extract the full name listed.
#             Extracted Value:

#         -   Primary Contact Email: The email address of the primary contact person.
#             Reasoning: Locate the section labeled "Primary Contact." Extract the email address listed.
#             Extracted Value:

#         -   Number of Employees: The total number of employees reported by the client.
#             Reasoning: Identify the section labeled "Number of Employees" and extract the numerical value listed.
#             Extracted Value:

#         -   Annual Revenue: The client's annual revenue for the most recent year.  Include the currency (e.g., USD).
#             Reasoning:  Identify the section labeled "Annual Revenue (Most Recent Year)" and extract the value. Include the currency.
#             Extracted Value:

#         -   UBOs (Ultimate Beneficial Owners): A list of the names and ownership percentages of each Ultimate Beneficial Owner (UBO).  Format the output as a comma-separated list of "Name (Percentage)".
#             Reasoning: Locate the section labeled "Ultimate Beneficial Owner(s) (UBOs)".  Extract the name and ownership percentage for each UBO.
#             Extracted Value:

#         -   Primary Bank (Currently): The name of the bank where the client currently holds their primary account.
#             Reasoning: Locate the section labeled "Primary Bank (Currently)" and extract the name of the bank.
#             Extracted Value:

#         -   Services Requested: A comma separated list of the banking services requested by the client (e.g., Checking Account, Online Banking).
#             Reasoning: Locate the section labeled "Banking Services Requested". Extract the names of all listed services.
#             Extracted Value:

#         Important Notes:

#         *   Do NOT attempt to access or extract information from external documents mentioned in the text (e.g., "Audited Financial Statements provided separately"). Focus ONLY on the information provided WITHIN the document text.
#         *   The information about UBO ownership is crucial for KYC compliance. If the information cannot be found for each UBO, respond with "Value not found".
#         *   Pay close attention to the formatting requirements for each extracted value.
#         """

#         # 3. Call Gemini API
#         gemini_output = extract_key_value_pairs(document_text, prompt_template)

#         if gemini_output:
#             print("Raw Gemini Output:\n", gemini_output)

#             # 4. Post-process the output
#             kv_pairs = post_process_kv_pairs_cot(gemini_output)  # Use the CoT post-processing

#             print("\nExtracted Key-Value Pairs (with Reasoning):", kv_pairs)
#             # for key, data in kv_pairs.items():
#             #     print(f"Key: {key}")
#             #     print(f"  Value: {data['value']}")
#             #     print(f"  Reasoning:\n{data['reasoning']}")  # Print the reasoning
#             #     print("-" * 20)
#             output_filename = "extracted_data.json"
#             with open(output_filename, "w") as outfile:
#                 json.dump(kv_pairs, outfile, indent=4)

#         else:
#             print("Key-value extraction failed.")
#     else:
#         print("Failed to read the document.")



def read_word_document(file_path):
    """Reads text from a Word document.  Uses python-docx library."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading document with docx: {e}")
        return None


def read_word_document_unstructured(file_path):
    """Reads text from a Word document using the unstructured library."""
    try:
        elements = partition(filename=file_path)
        full_text = "\n".join([str(el.text) for el in elements])
        return full_text
    except Exception as e:
        print(f"Error reading document with unstructured: {e}")
        return None


def extract_key_value_pairs(text, prompt_template):
    """Extracts key-value pairs using Gemini API."""
    prompt = prompt_template.format(text=text)

    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error during Gemini API call: {e}")
        return None


def post_process_kv_pairs(gemini_output):
    """
    Post-processes the raw Gemini output to extract clean key-value pairs.
    This function is crucial for handling inconsistencies in Gemini's output format.
    """
    key_value_pairs = {}
    lines = gemini_output.split("\n")

    for line in lines:
        # Split the line by a colon, but only the first one
        parts = line.split(":", 1)

        if len(parts) == 2:  # Ensure there's a key and a value
            key = parts[0].strip()
            value = parts[1].strip()
            key_value_pairs[key] = value
        else:
            #Handle lines that don't split properly (e.g., just a key or just a value)
            print(f"Skipping malformed line: {line}")

    return key_value_pairs


if __name__ == "__main__":
    file_path = "sample.docx"  # Replace with your Word document path

    # 1. Read the Document
    #document_text = read_word_document(file_path) # using python-docx

    document_text = read_word_document_unstructured(file_path) #Using unstructured


    if document_text:
        # 2. Define the Prompt (Crucial)
        prompt_template = """
        You are an expert financial onboarding specialist. Your task is to extract key information from a commercial client onboarding document.

        Here's the document content you will be reviewing:

        Text:
        {text}

        Based on the document content provided, extract the following information.  If a value is not found or cannot be reliably determined from the provided text, explicitly state "Value not found".

        Provide the Extracted Value ONLY after the key.

        Specifically, extract the following:

        Client Legal Name:
        Doing Business As :
        Industry :
        Date of Incorporation :
        Client ID:
        Primary Contact Name:
        Primary Contact Email:
        Number of Employees:
        Annual Revenue:
        UBOs (Ultimate Beneficial Owners):
        Primary Bank (Currently):
        Services Requested:

        Important Notes:

        *   Do NOT attempt to access or extract information from external documents mentioned in the text (e.g., "Audited Financial Statements provided separately"). Focus ONLY on the information provided WITHIN the document text.
        *   The information about UBO ownership is crucial for KYC compliance. If the information cannot be found for each UBO, respond with "Value not found".
        *   Pay close attention to the formatting requirements for each extracted value.
        """

        # 3. Call Gemini API
        gemini_output = extract_key_value_pairs(document_text, prompt_template)

        if gemini_output:
            print("Raw Gemini Output:\n", gemini_output)

            # 4. Post-process the output
            kv_pairs = post_process_kv_pairs(gemini_output)

            print("\nExtracted Key-Value Pairs:")
            for key, value in kv_pairs.items():
                print(f"{key}: {value}")

            # 5. Save to JSON
            output_filename = "extracted_data.json"
            with open(output_filename, "w") as outfile:
                json.dump(kv_pairs, outfile, indent=4)  # Save as JSON with indentation

            print(f"\nExtracted data saved to {output_filename}")

        else:
            print("Key-value extraction failed.")
    else:
        print("Failed to read the document.")